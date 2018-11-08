import os, re, json, pypandoc
from io import BytesIO
from zipfile import ZipFile
from django.http import HttpResponse
from django.utils.timezone import datetime
from django.conf import settings
from tempfile import NamedTemporaryFile
from docxtpl import DocxTemplate, RichText
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from lxml import etree
from .models import Plans, UmkArticles, UmkData, Competence, User
from .core_funcs import previous_and_next_disciplines_from_umk

def is_int(n):
    return not n%1

def upp(val, condition):
    if len(val)==condition:
        res = True if int(val[0])>0 else False
    else:
        res = False
    return res

def isEmptyVal(value):
    try:
        val = float(value)
        if is_int(val):
            val = int(val)
        return val if val>0 else "-"
    except:
        return "-"

def isEmptyValOrStr(value1,value2,value3):
    return "{0}/{1}/{2}".format(isEmptyVal(value1), isEmptyVal(value2), isEmptyVal(value3))

def remove_htmk_tags(val):
    return "{0}".format(re.compile(r'<.*?>').sub('',val))

def html_to_docx(html_text, doc_tpl): #конвертирование html кода в docx
    file_object = NamedTemporaryFile(suffix='.docx')# create temp file
    pypandoc.convert(html_text, 'docx', format='html', outputfile=file_object.name) # generate it using pandoc
    subdocx = DocxTemplate(file_object.name)#open docx file
    subdocx._part = doc_tpl.docx._part
    if subdocx._element.body.sectPr is not None:
        subdocx._element.body.remove(subdocx._element.body.sectPr)
    xml = re.sub(r'</?w:body[^>]*>', '', etree.tostring(subdocx._element.body,encoding='unicode', pretty_print=False))
    return xml


def num_is_empty(str):
    if len(str)>0:
        v = str.split("/")
        res = isEmptyValOrStr(v[0],v[1],v[2]) if len(v)>1 else isEmptyVal(v[0])
    else:
        res = "-"
    return res


def get_course(s):

    tmp = []
    sem = s.split(",")

    for n in sem:
        l = int(n)
        [course, l] = divmod(l,2)
        if(l !=0):
            course += 1
        tmp.append(str(course))

    res = []
    for i in tmp:
        if i not in res:
            res.append(i)

    return ",".join(res)

#-------------------------------------------------------------------------------------------

def get_competens(plan):
    res = []
    list = plan.comps.split(" ")

    for v in list:
        if v:
            cmps = Competence.objects.filter(name=v, direction=plan.direction) #, training_program__contains=plan.training_program.split(" ")[0]
            for c in cmps:
                print("comp: " + v + " " + c.training_program + " search: " + plan.training_program)
                if c.training_program == plan.training_program:
                    res.append({'name': c.name, 'full_content': c.full_content, 'student_known': c.should_know, 'student_can': c.should_able, 'student_own': c.should_master})


    return res

def get_zaf_kaf(deparmt):
    user = User.objects.filter(deparmt=deparmt)
    res = {"name": '', "position": ''}

    for u in user:
        if u.position == "io_zaf_kaf" or u.position == "zaf_kaf":
            res = {'name': "{0} {1}. {2}.".format(u.last_name, u.first_name[0], u.patronymic[0]),   'position': u.get_position_display()}

    return res

def get_predsedatel_spn(deparmt):
    user = User.objects.filter(deparmt=deparmt)
    res = ''

    for u in user:
        if u.position == "predsedatel_spn":
            res = "{0} {1}. {2}.".format(u.last_name, u.first_name[0], u.patronymic[0])

    return res

def required_reconcil(umkcreator, plans):
    rt = RichText()
    if umkcreator.deparmt != plans[0].direction.deparmt:
        rt.add("СОГЛАСОВАНО:\n")
        zaf_kaf = get_zaf_kaf(plans[0].direction.deparmt)
        rt.add("{0} выпускающей кафедрой ________________ {1}\n".format(zaf_kaf['position'].replace("каф.",""),zaf_kaf['name']))
        rt.add("«____»___________20__г.\n")

    return rt

def get_OPOP_of_discipline(plans):
    code_OPOP = plans[0].code_OPOP

    if len(re.findall("Б", code_OPOP))==2:
        res = 'базовой части {0}'.format(code_OPOP)
    elif len(re.findall("Б", code_OPOP))==1 and len(re.findall("В", code_OPOP))==1:
        res = 'вариативной части {0}'.format(code_OPOP)
    elif len(re.findall("Б", code_OPOP))==1 and len(re.findall("В", code_OPOP))==2:
        res = 'дисциплине по выбору {0}'.format(code_OPOP)
    else:
        res = "не заполнено"

    return res

def get_placeInStructOPOP(plans, discipline, doc_tpl):
    prev_next_discip = previous_and_next_disciplines_from_umk(plans[0],Plans)
    res = ["",""]
    if prev_next_discip['previous_disciplines']:
        res[0] = "Для полного усвоения данной дисциплины обучающиеся должны знать следующие дисциплины: "
        for d in prev_next_discip['previous_disciplines']:
            res[0] += " {0} - {1};".format(d.code_OPOP, d.discipline.name)

        if res[0].endswith(";"):
            res[0] = res[0][0:len(res[0])-1]
            res[0] +="."

    if prev_next_discip['next_disciplines']:
        res[1] += "Знания по дисциплине \"{0}\" необходимы обучающимся данного направления для усвоения знаний по следующим дисциплинам: ".format(discipline)
        for d in prev_next_discip['next_disciplines']:
            res[1] += " {0} - {1};".format(d.code_OPOP, d.discipline.name)

        if res[1].endswith(";"):
            res[1] = res[1][0:len(res[1])-1]
            res[1] +="."

    sd = doc_tpl.new_subdoc()
    st1 = sd.add_paragraph(res[0])
    st1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    st2 = sd.add_paragraph(res[1])
    st2.paragraph_format.first_line_indent = Cm(1.0) #абзацный отступ
    st2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY #выравнивание по центру

    return sd

def get_table_contentsection(umk_data):
    res = []
    js = json.JSONDecoder()
    sec = js.decode(umk_data.contentOfSections)

    for item in sec:
      res.append({'id':item[0], 'name': item[1], 'content': item[2]})

    return res

def get_table_interdisciplinary_relations(umk_data):
    res = []
    js = json.JSONDecoder()
    sec = js.decode(umk_data.interdiscipRelations)

    for item in sec:
        res.append({'id': item[0], 'name': item[1], 'num_sec': item[2]})

    return res


def get_table_sections_hours(umk_data):
    res = []
    js = json.JSONDecoder()
    list = js.decode(umk_data.table_sections_hour)

    for item in list:
        res.append({'id': item[0],    'name': item[1], 'lec_h': num_is_empty(item[2]), 'prakt_h': num_is_empty(item[3]), 'lab_h': num_is_empty(item[4]),
                    'sem_h': num_is_empty(item[5]), 'cpc_h': num_is_empty(item[6]), 'total_h': num_is_empty(item[7]), 'inter_h': num_is_empty(item[8])})

    return res


def get_table_lections_hours(umk_data):
    res = []
    js = json.JSONDecoder()
    list = js.decode(umk_data.table_lectures_hour)

    for item in list:
        res.append({'id_sec': item[0], 'id_theme': item[1], 'name_lec': item[2], 'hours': item[3], 'comps': item[4], 'methods': item[5]})
    return res


def get_table_labs_prakt(data,doc_tpl,type):
    js = json.JSONDecoder()
    sd = doc_tpl.new_subdoc()

    if data:
        list = js.decode(data)
        if list[0][2] !='':
            table = sd.add_table(rows=1, cols=6,style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '№ п/п'
            hdr_cells[1].text = '№ темы'
            hdr_cells[2].text = 'Темы {0} работ'.format(type)
            hdr_cells[3].text = 'Трудоемкость (час.)'
            hdr_cells[4].text = 'Формируемые компетенции'
            hdr_cells[5].text = 'Методы преподавания'

            for item in list:
                row_cells = table.add_row().cells
                row_cells[0].text = item[0]
                row_cells[1].text = item[1]
                row_cells[2].text = item[2]
                row_cells[3].text = num_is_empty(item[3])
                row_cells[4].text = item[4]
                row_cells[5].text = item[5]
        else:
            sd.add_paragraph('Учебным планом {0} работ не предусмотрено'.format(type))
    else:
        sd.add_paragraph('Учебным планом {0} работ не предусмотрено'.format(type))

    return sd

def get_table_samost_hours(umk_data):
    res = []
    js = json.JSONDecoder()
    if umk_data.table_samost_hour:
        list = js.decode(umk_data.table_samost_hour)

        for item in list:
            res.append({'id': item[0], 'num_sec': item[1], 'name_theme': item[2], 'hours': num_is_empty(item[3]), 'kontrol': item[4],
                        'comps': item[5]})
    return res

def get_hour_kursovaya_work_or_project(umk_data): #возвращает количество часов отведенных на курсовой проект или работу
    res = "-/-/-"
    js = json.JSONDecoder()
    if umk_data.table_samost_hour:
        list = js.decode(umk_data.table_samost_hour)

        for item in list:
            if re.findall('Выполнение курсового проекта / курсовой работы', item[2]):
                res = num_is_empty(item[3])
    return res

def get_hour_raschetno_graph_work(umk_data):#возвращает количество часов отведенных на расчетно-графическую работу
    res = "-/-/-"
    js = json.JSONDecoder()
    if umk_data.table_samost_hour:
        list = js.decode(umk_data.table_samost_hour)

        for item in list:
            if re.findall('Выполнение расчетно-графических домашних работ', item[2]):
                res = num_is_empty(item[3])
    return res


def get_table_literature(umkdata):
    res = []
    js = json.JSONDecoder()
    if umkdata.table_literature:
        list = js.decode(umkdata.table_literature)

        for item in list:
            res.append({'main_add': item[0], 'name': item[1], 'year': item[2], 'type': item[3], 'vid': item[4],
                        'count': item[5], 'kontingent': item[6], 'obespechennost': item[7], 'place': item[8], 'electr_variant': item[9]})
    return res

def get_kontrolnya(plans): #только для заочной формы обучения
    res = ''
    if plans[1].kontrolnaya_work == plans[2].kontrolnaya_work and  ( (plans[1].trudoemkost_all - plans[1].hours_samost_work_sum) == (plans[2].trudoemkost_all - plans[2].hours_samost_work_sum)):
        res = 'Для заочной формы обучения предусмотрена контрольная работа объемом {0} часов в {1} семестре. '.format(plans[1].trudoemkost_all - plans[1].hours_samost_work_sum, plans[1].kontrolnaya_work)
    else:
        for i in range(1,3):
            srok = re.findall(r"\d\.\d|\d+", plans[i].get_training_form_display())[0]
            if plans[i].kontrolnaya_work:
                res += 'Для заочной ({0} {1}) формы обучения предусмотрена контрольная работа объемом {2} часов в {3} семестре. '.format(srok, "лет" if float(srok)>=5 else "года",
                                                                                                                                         plans[i].trudoemkost_all - plans[i].hours_samost_work_sum,
                                                                                                                                         plans[i].kontrolnaya_work)

    return res

def get_table_rating_day(doc_tpl, plan, umkdata):#Оценка результатов освоения учебной дисциплины для очной формы
    js = json.JSONDecoder()
    datafortable = []
    tmp = []
    id = 0

    sd = doc_tpl.new_subdoc()
    if umkdata.table_rating_ochka:
        sd.add_paragraph('Рейтинговая система оценки по дисциплине «{0}» для обучающихся направления {1} очной формы обучения'.format(plan.discipline.name, plan.get_direction()))
        table1 = sd.add_table(rows=1, cols=4, style='Table Grid')
        hdr_cells = table1.rows[0].cells
        hdr_cells[0].text = '1-ый срок предоставления результатов текущего контроля'
        hdr_cells[1].text = '2-ый срок предоставления результатов текущего контроля'
        hdr_cells[2].text = '3-ый срок предоставления результатов текущего контроля'
        hdr_cells[3].text = 'Итоговый тест'

        for item in js.decode(umkdata.table_rating_ochka):
            if re.search("Итого за \d-ую аттестацию", item[1]):
                tmp.append(item[2].split("-")[1] if re.search("0-\d+", item[2]) else item[2])
                id += 1
            if id == 3:
                datafortable.append(tmp)
                tmp = []
                id = 0

        for item in datafortable:
            row_cells = table1.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
            row_cells[2].text = item[2]
            row_cells[3].text = '100'

        sd.add_paragraph(" ")
        table = sd.add_table(rows=1, cols=4, style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'Виды контрольных мероприятий'
        hdr_cells[2].text = 'Баллы'
        hdr_cells[3].text = '№ недели'

        for item in js.decode(umkdata.table_rating_ochka):
            row_cells = table.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
            row_cells[2].text = item[2]
            row_cells[3].text = item[3]

    return sd

def get_table_rating_night(doc_tpl, umkdata):#Оценка результатов освоения учебной дисциплины для заочной формы
    js = json.JSONDecoder()
    sd = doc_tpl.new_subdoc()

    if umkdata.table_rating_zaochka:
        table = sd.add_table(rows=1, cols=3, style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'Виды контрольных мероприятий'
        hdr_cells[2].text = 'Баллы'

        for item in js.decode(umkdata.table_rating_zaochka):
            row_cells = table.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
            row_cells[2].text = item[2]

    return sd
#####################################################################################################################
#
#
#                       Формирование рабочей программы и аннотации
#
#
######################################################################################################################
def render_context(id, doc_tpl):
    umk = UmkArticles.objects.get(id=id)
    umkdata = UmkData.objects.get(umk_id=id)
    plans = [
        Plans.objects.get(id=umk.plan_ochka),
        Plans.objects.get(id=umk.plan_z),
        Plans.objects.get(id=umk.plan_zu),
    ]
    dateprikaz = plans[0].get_direction_date_prikaz()
    kursovya_work_project_sem = [re.findall("\d+",plan.kursovya_work_project)[0] for plan in plans if re.findall("\d+",plan.kursovya_work_project)] #номер семестра на курсовой проект

    context = {'ministerstvo': umk.creator.deparmt.units.univer.ministerstvo,
               'UNIVERCITY': umk.creator.deparmt.units.univer.name,
               'Units': umk.creator.deparmt.units.name,
               'kafedra': umk.creator.deparmt.name,
               'predsedatel_spn': get_predsedatel_spn(plans[0].direction.deparmt),
               'zav_kaf': get_zaf_kaf(umk.creator.deparmt),
               'author': {'name': umk.creator.get_fullname(),
                          'position': umk.creator.get_position_display(),
                          'academic_degree': umk.creator.get_science_stepen_display(),
                          'rank': ", {0}".format(umk.creator.get_science_zvanie_display()) if umk.creator.science_zvanie!='no' else ""},
               'Discipline': plans[0].get_discipline(),
               'direction': plans[0].get_direction(),
               'Profiles': plans[0].get_profiles(),
               'qualification': plans[0].get_qualif_display(),
               'program_traning': "{0} {1}а".format(plans[0].get_training_program_display(),plans[0].get_qualif_display()),
               'cources': "{0}/{1}/{2}".format(get_course(plans[0].get_semestrs()),get_course(plans[1].get_semestrs()),get_course(plans[2].get_semestrs())),                                                   #нужно рассчитать изходя из семестров
               'semestrs': "{0}/{1}/{2}".format(plans[0].get_semestrs(), plans[1].get_semestrs(),
                                                plans[2].get_semestrs()),
               'audit_total_hours': "{0}/{1}/{2}".format(plans[0].hours_audit_work_sum, plans[1].hours_audit_work_sum,
                                                         plans[2].hours_audit_work_sum),
               'lectures_total_hours': isEmptyValOrStr(plans[0].hours_lectures, plans[1].hours_lectures,
                                                       plans[2].hours_lectures),
               'prakt_total_hours': isEmptyValOrStr(plans[0].hours_pract, plans[1].hours_pract, plans[2].hours_pract),
               'labs_total_hours': isEmptyValOrStr(plans[0].hours_labs, plans[1].hours_labs, plans[2].hours_labs),
               'samost_total_hours': "{0}/{1}/{2}".format(plans[0].hours_samost_work_sum,
                                                          plans[1].hours_samost_work_sum,
                                                          plans[2].hours_samost_work_sum),
               'kursovya_num_semestr': isEmptyValOrStr(kursovya_work_project_sem[0],kursovya_work_project_sem[1],kursovya_work_project_sem[2]) if kursovya_work_project_sem else "-",
               'kursovya_hours': get_hour_kursovaya_work_or_project(umkdata),
               'raschotno_graph_work_semests': "-",
               'raschotno_graph_work_hours': "-",
               'zanyatiya_in_interaktiv_hours': "{0}".format(plans[0].zanatiya_in_interak_forms_hours),
               'zachot_semestrs': "{0}/{1}/{2}".format(plans[0].zachot_semestr if len(plans[0].zachot_semestr.split(','))>1 else isEmptyVal(plans[0].zachot_semestr),
                                                       plans[1].zachot_semestr if len(plans[1].zachot_semestr.split(',')) > 1 else isEmptyVal(plans[1].zachot_semestr),
                                                       plans[2].zachot_semestr if len(plans[2].zachot_semestr.split(',')) > 1 else isEmptyVal(plans[2].zachot_semestr)),
               'exam_semestrs': "{0}/{1}/{2}".format(plans[0].exam_semestr if len(plans[0].exam_semestr.split(','))>1 else isEmptyVal(plans[0].exam_semestr),
                                                       plans[1].exam_semestr if len(plans[1].exam_semestr.split(',')) > 1 else isEmptyVal(plans[1].exam_semestr),
                                                       plans[2].exam_semestr if len(plans[2].exam_semestr.split(',')) > 1 else isEmptyVal(plans[2].exam_semestr)),
               'trudoemkost_all': "{0}".format(plans[0].trudoemkost_all),
               'zachot_edinic': "{0}".format(plans[0].trudoemkost_zachot_edinic),
               'prikaz_number': dateprikaz[0],
               'prikaz_date': dateprikaz[1].strftime("%d.%m.%Y"),
               'year': datetime.now().strftime("%Y"),
               'req_reconciliation': required_reconcil(umk.creator, plans),
               # -------------------------------------------------------------------------------------------------------
               'discipline_aims':  html_to_docx(umkdata.aim, doc_tpl),
               'discipline_tasks': html_to_docx(umkdata.tasks, doc_tpl),
               'code_OPOP': get_OPOP_of_discipline(plans),
               'discipline_place': get_placeInStructOPOP(plans, plans[0].get_discipline(), doc_tpl),
               'tbl_comps': get_competens(plans[0]),
               'tbl_contentsections_and_theme': get_table_contentsection(umkdata),
               'tbl_interdisciplinary_rel': get_table_interdisciplinary_relations(umkdata),
               'tbl_section_hours': get_table_sections_hours(umkdata),
               'tbl_lections_hours': get_table_lections_hours(umkdata),
               'tbl_prakt_hours': get_table_labs_prakt(umkdata.table_prakt_hour, doc_tpl, 'практических'),
               'tbl_labs_hours': get_table_labs_prakt(umkdata.table_laborat_hour, doc_tpl, 'лабораторных'),
               'tbl_samost_hours': get_table_samost_hours(umkdata),
               'tbl_liter': get_table_literature(umkdata),
               'samost_total': "{0}/{1}/{2}".format(plans[0].hours_samost_work_sum,plans[1].hours_samost_work_sum,plans[2].hours_samost_work_sum),
               'samost_total_without_prepod': "{0}/{1}/{2}".format(plans[0].hours_samost_wo_lec,plans[1].hours_samost_work_sum,plans[2].hours_samost_work_sum),
               'samost_total_with_student': "{0}/-/-".format(plans[0].hours_samost_w_lec_w_stud),
               'samost_total_with_group': "{0}/-/-".format(plans[0].hours_samost_w_lec_w_group),
               'theme_kursovii_work': umkdata.theme_kursovih_rabot if len(umkdata.theme_kursovih_rabot)>1 else 'Учебным планом выполнение курсовых работ не предусмотрено.',
               'materialno_texnicheskoe_obespechenie': html_to_docx(umkdata.material_teh_obespech_dicip, doc_tpl),
               'database_info_system': html_to_docx(umkdata.database_info_system, doc_tpl),
               #-------------------------------------------Рейтинг------------------------------------------------------
               'rating_day': get_table_rating_day(doc_tpl,plans[0], umkdata),
               'rating_night':get_table_rating_night(doc_tpl, umkdata),
               #-----------------------------------------ЛИТЕРАТУРА-----------------------------------------------------
               'crs_och': get_course(plans[0].get_semestrs()),
               'crs_z': get_course(plans[1].get_semestrs()),
               'crs_zu': get_course(plans[2].get_semestrs()),
               'smr_och': plans[0].get_semestrs(),
               'smr_z': plans[1].get_semestrs(),
               'smr_zu': plans[2].get_semestrs()
               }
    return context

#    Создание/Генерирование аннотации
def render_context_annotation(id, doc_tpl):
    umkdata = UmkData.objects.get(umk_id=id)
    umk = umkdata.umk_id
    plans = [
        Plans.objects.get(id=umk.plan_ochka),
        Plans.objects.get(id=umk.plan_z),
        Plans.objects.get(id=umk.plan_zu),
    ]

    student_known = set()
    student_can = set()
    student_own = set()


    for v in plans[0].comps.split(" "):
        if v:
            cmps = Competence.objects.filter(name=v, direction=plans[0].direction)
            for c in cmps:
                print("comp: " + v + " " + c.training_program + " search: " + plans[0].training_program)
                if c.training_program == plans[0].training_program:
                    student_known.add(c.should_know)
                    student_can.add(c.should_able)
                    student_own.add(c.should_master)

    context = {'zav_kaf': get_zaf_kaf(umk.creator.deparmt),
               'author': {'name': umk.creator.get_fullname(),
                          'position': umk.creator.get_position_display(),
                          'academic_degree': umk.creator.get_science_stepen_display(),
                          'rank': ", {0}".format(umk.creator.get_science_zvanie_display()) if umk.creator.science_zvanie!='no' else ""},
               'discipline': plans[0].get_discipline(),
               'direction': plans[0].get_direction(),
               'audit_total_hours': "{0}/{1}/{2}".format(plans[0].hours_audit_work_sum, plans[1].hours_audit_work_sum,
                                                         plans[2].hours_audit_work_sum),
               'samost_total_hours': "{0}/{1}/{2}".format(plans[0].hours_samost_work_sum,
                                                          plans[1].hours_samost_work_sum,
                                                          plans[2].hours_samost_work_sum),
               'zachot_semestrs': "{0}/{1}/{2}".format(plans[0].zachot_semestr if len(plans[0].zachot_semestr.split(','))>1 else isEmptyVal(plans[0].zachot_semestr),
                                                       plans[1].zachot_semestr if len(plans[1].zachot_semestr.split(',')) > 1 else isEmptyVal(plans[1].zachot_semestr),
                                                       plans[2].zachot_semestr if len(plans[2].zachot_semestr.split(',')) > 1 else isEmptyVal(plans[2].zachot_semestr)),
               'exam_semestrs': "{0}/{1}/{2}".format(plans[0].exam_semestr if len(plans[0].exam_semestr.split(','))>1 else isEmptyVal(plans[0].exam_semestr),
                                                       plans[1].exam_semestr if len(plans[1].exam_semestr.split(',')) > 1 else isEmptyVal(plans[1].exam_semestr),
                                                       plans[2].exam_semestr if len(plans[2].exam_semestr.split(',')) > 1 else isEmptyVal(plans[2].exam_semestr)),
               'trudoemkost_all': "{0}".format(plans[0].trudoemkost_all),
               'year_nabor': plans[0].year,
               # -------------------------------------------------------------------------------------------------------
               'discipline_aims':  html_to_docx(umkdata.aim, doc_tpl),
               'code_OPOP': get_OPOP_of_discipline(plans),
               'discipline_place': get_placeInStructOPOP(plans, plans[0].get_discipline(), doc_tpl),
               'comps':plans[0].comps,
               'student_known': " ".join(student_known),
               'student_can': " ".join(student_can),
               'student_own': " ".join(student_own)
               }
    return context


def generation_docx(id):#формирование только одной рабочей программы и аннотации
    #doc = DocxTemplate(os.path.join(os.path.join(settings.BASE_DIR, "static/doc_templ"), 'template.docx'))
    #doc.render(render_context(id,doc))

    #tmpfile = BytesIO()
    #doc.save(tmpfile)
    #res = HttpResponse(tmpfile.getvalue(),
                       #content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    #res['Content-Disposition'] = 'attachment; filename=result.docx'
    #res['Content-Length'] = tmpfile.tell()

    umkname = UmkArticles.objects.get(id=int(id)).get_short_name() #имя дисциплины и профиль

    tmpfile = BytesIO()
    with ZipFile(tmpfile, 'w') as myzip:
        ####создание рабочей программы
        doc = DocxTemplate(os.path.join(os.path.join(settings.BASE_DIR, "static/doc_templ"), 'template.docx'))
        doc.render(render_context(int(id), doc))
        rp_file_object = NamedTemporaryFile(prefix="{0}-".format(umkname),suffix='.docx')  # create temp file
        doc.save(rp_file_object.name)
        myzip.write(rp_file_object.name)
        rp_file_object.close()

        # создание аннотации
        ann_doc = DocxTemplate(os.path.join(os.path.join(settings.BASE_DIR, "static/doc_templ"), 'template_annotation.docx'))
        ann_doc.render(render_context_annotation(id, ann_doc))

        ann_file_object = NamedTemporaryFile(prefix="{0}-аннотация-".format(umkname),suffix='.docx')  # create temp file
        ann_doc.save(ann_file_object.name)
        myzip.write(ann_file_object.name)
        ann_file_object.close()

    res = HttpResponse(tmpfile.getvalue(), content_type='application/zip')
    res['Content-Disposition'] = 'attachment; filename=result-{0}.zip'.format(datetime.today().strftime("%Y-%m-%d"))
    res['Content-Length'] = tmpfile.tell()
    return res

def generation_docx_achive(id_list):
    tmpfile = BytesIO()
    with ZipFile(tmpfile, 'w') as myzip:
        for id in id_list:
            ####создание рабочей программы
            doc = DocxTemplate(os.path.join(os.path.join(settings.BASE_DIR, "static/doc_templ"), 'template.docx'))
            doc.render(render_context(int(id),doc))
            prefix = "{0}-".format(UmkArticles.objects.get(id=int(id)).get_short_name())
            rp_file_object = NamedTemporaryFile(prefix=prefix, suffix='.docx')  # create temp file
            doc.save(rp_file_object.name)
            myzip.write(rp_file_object.name)
            rp_file_object.close()

            #создание аннотации
            ann_doc = DocxTemplate(os.path.join(os.path.join(settings.BASE_DIR, "static/doc_templ"), 'template_annotation.docx'))
            ann_doc.render(render_context_annotation(id, ann_doc))

            ann_file_object = NamedTemporaryFile(prefix="{0}-аннотация-".format(UmkArticles.objects.get(id=int(id)).get_short_name()), suffix='.docx')  # create temp file
            ann_doc.save(ann_file_object.name)
            myzip.write(ann_file_object.name)
            ann_file_object.close()

    res = HttpResponse(tmpfile.getvalue(), content_type='application/zip')
    res['Content-Disposition'] = 'attachment; filename=result-{0}.zip'.format(datetime.today().strftime("%Y-%m-%d"))
    res['Content-Length'] = tmpfile.tell()
    return res

