import re,json, pypandoc
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from lxml import etree
from tempfile import NamedTemporaryFile
from docxtpl import DocxTemplate
from .models import Plans, Competence, User
from .core_funcs import previous_and_next_disciplines_from_umk

def is_int(n):
    return not n%1

def upp(val, condition):
    if len(val)==condition:
        res = True if int(val[0])>0 else False
    else:
        res = False
    return res

def isEmptyVal(value):
    try:
        val = float(value)
        if is_int(val):
            val = int(val)
        return val if val>0 else "-"
    except:
        return "-"

def isEmptyValOrStr(value1,value2,value3):
    return "{0}/{1}/{2}".format(isEmptyVal(value1), isEmptyVal(value2), isEmptyVal(value3))

def remove_htmk_tags(val):
    return "{0}".format(re.compile(r'<.*?>').sub('',val))

def html_to_docx(html_text, doc_tpl): #конвертирование html кода в docx
    file_object = NamedTemporaryFile(suffix='.docx')# create temp file
    pypandoc.convert(html_text, 'docx', format='html', outputfile=file_object.name) # generate it using pandoc
    subdocx = DocxTemplate(file_object.name)#open docx file
    subdocx._part = doc_tpl.docx._part
    if subdocx._element.body.sectPr is not None:
        subdocx._element.body.remove(subdocx._element.body.sectPr)
    xml = re.sub(r'</?w:body[^>]*>', '', subdocx.get_xml())#etree.tostring(subdocx._element.body,encoding='unicode', pretty_print=False))
    return xml


def num_is_empty(str):
    if len(str)>0:
        v = str.split("/")
        res = isEmptyValOrStr(v[0],v[1],v[2]) if len(v)>1 else isEmptyVal(v[0])
    else:
        res = "-"
    return res

#----------------------------------------------------------------------------
def get_course(s):

    tmp = []
    sem = s.split(",")

    for n in sem:
        l = int(n)
        [course, l] = divmod(l,2)
        if(l !=0):
            course += 1
        tmp.append(str(course))

    res = []
    for i in tmp:
        if i not in res:
            res.append(i)

    return ",".join(res)

def get_competens(plan):
    res = []
    list = plan.comps.split(" ")

    for v in list:
        if v:
            cmps = Competence.objects.filter(name=v, direction=plan.direction)
            for c in cmps:
                print("comp: " + v + " " + c.training_program + " search: " + plan.training_program)
                if c.training_program == plan.training_program:
                    res.append({'name': c.name, 'full_content': c.full_content, 'student_known': c.should_know, 'student_can': c.should_able, 'student_own': c.should_master,
                                'indicators_know': c.indicators_know, 'indicators_can': c.indicators_can, 'indicators_own': c.indicators_own})


    return res

def get_zaf_kaf(deparmt):#получаем ФИО зав кафедрой
    user = User.objects.filter(deparmt=deparmt)
    res = {"name": '', "position": ''}

    for u in user:
        if u.position == "io_zaf_kaf" or u.position == "zaf_kaf":
            res = {'name': "{0} {1}. {2}.".format(u.last_name, u.first_name[0], u.patronymic[0]),   'position': u.get_position_display()}

    return res

def get_predsedatel_spn(deparmt):#председатель СПН
    user = User.objects.filter(deparmt=deparmt)
    res = ''
    for u in user:
        if u.position == "predsedatel_spn":
            res = "{0} {1}. {2}.".format(u.last_name, u.first_name[0], u.patronymic[0])

    return res

def required_reconcil(umkcreator, plans):#формирование надписи согласовано
    zaf_kaf = get_zaf_kaf(plans[0].direction.deparmt)
    return {'position': "{0} выпускающей кафедрой".format(zaf_kaf['position'].replace("каф.","")), 'name': zaf_kaf['name']}

def get_OPOP_of_discipline(plans):
    code_OPOP = plans[0].code_OPOP

    if len(re.findall("Б", code_OPOP))==2:
        res = 'базовой части {0}'.format(code_OPOP)
    elif len(re.findall("Б", code_OPOP))==1 and len(re.findall("В", code_OPOP))==1:
        res = 'вариативной части {0}'.format(code_OPOP)
    elif len(re.findall("Б", code_OPOP))==1 and len(re.findall("В", code_OPOP))==2:
        res = 'дисциплине по выбору {0}'.format(code_OPOP)
    else:
        res = "не заполнено"

    return res

def get_placeInStructOPOP(plans, discipline, doc_tpl):
    prev_next_discip = previous_and_next_disciplines_from_umk(plans[0],Plans)
    res = ["",""]
    if prev_next_discip['previous_disciplines']:
        res[0] = "Для полного усвоения данной дисциплины обучающиеся должны знать следующие дисциплины: "
        for d in prev_next_discip['previous_disciplines']:
            res[0] += " {0} - {1};".format(d.code_OPOP, d.discipline.name)

        if res[0].endswith(";"):
            res[0] = res[0][0:len(res[0])-1]
            res[0] +="."

    if prev_next_discip['next_disciplines']:
        res[1] += "Знания по дисциплине \"{0}\" необходимы обучающимся данного направления для усвоения знаний по следующим дисциплинам: ".format(discipline)
        for d in prev_next_discip['next_disciplines']:
            res[1] += " {0} - {1};".format(d.code_OPOP, d.discipline.name)

        if res[1].endswith(";"):
            res[1] = res[1][0:len(res[1])-1]
            res[1] +="."

    sd = doc_tpl.new_subdoc()
    st1 = sd.add_paragraph(res[0])
    st1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    st2 = sd.add_paragraph(res[1])
    st2.paragraph_format.first_line_indent = Cm(1.0) #абзацный отступ
    st2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY #выравнивание по центру

    return sd

def get_table_contentsection(umk_data):
    res = []
    js = json.JSONDecoder()
    if umk_data.contentOfSections:
        sec = js.decode(umk_data.contentOfSections)
        for item in sec:
          res.append({'id':item[0], 'name': item[1], 'content': item[2]})

    return res

def get_table_interdisciplinary_relations(umk_data):
    res = []
    js = json.JSONDecoder()
    if umk_data.interdiscipRelations:
        sec = js.decode(umk_data.interdiscipRelations)
        for item in sec:
            res.append({'id': item[0], 'name': item[1], 'num_sec': item[2]})

    return res

def get_table_sections_hours(umk_data):
    res = []
    js = json.JSONDecoder()
    list = js.decode(umk_data.table_sections_hour)

    for item in list:
        res.append({'id': item[0],    'name': item[1], 'lec_h': num_is_empty(item[2]), 'prakt_h': num_is_empty(item[3]), 'lab_h': num_is_empty(item[4]),
                    'sem_h': num_is_empty(item[5]), 'cpc_h': num_is_empty(item[6]), 'total_h': num_is_empty(item[7]), 'inter_h': num_is_empty(item[8])})

    return res

def get_table_lections_hours(umk_data):
    res = []
    js = json.JSONDecoder()
    list = js.decode(umk_data.table_lectures_hour)

    for item in list:
        res.append({'id_sec': item[0], 'id_theme': item[1], 'name_lec': item[2], 'hours': item[3], 'comps': item[4], 'methods': item[5]})
    return res


def get_table_labs_prakt(data,doc_tpl,type):
    js = json.JSONDecoder()
    sd = doc_tpl.new_subdoc()

    if data:
        list = js.decode(data)
        if list[0][2] !='':
            table = sd.add_table(rows=1, cols=6,style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '№ п/п'
            hdr_cells[1].text = '№ темы'
            hdr_cells[2].text = 'Темы {0} работ'.format(type)
            hdr_cells[3].text = 'Трудоемкость (час.)'
            hdr_cells[4].text = 'Формируемые компетенции'
            hdr_cells[5].text = 'Методы преподавания'

            for item in list:
                row_cells = table.add_row().cells
                row_cells[0].text = item[0]
                row_cells[1].text = item[1]
                row_cells[2].text = item[2]
                row_cells[3].text = num_is_empty(item[3])
                row_cells[4].text = item[4]
                row_cells[5].text = item[5]
        else:
            sd.add_paragraph('Учебным планом {0} работ не предусмотрено'.format(type))
    else:
        sd.add_paragraph('Учебным планом {0} работ не предусмотрено'.format(type))

    return sd

def get_table_samost_hours(umk_data):
    res = []
    js = json.JSONDecoder()
    if umk_data.table_samost_hour:
        list = js.decode(umk_data.table_samost_hour)

        for item in list:
            res.append({'id': item[0], 'num_sec': item[1], 'name_theme': item[2], 'hours': num_is_empty(item[3]), 'kontrol': item[4],
                        'comps': item[5]})
    return res

def get_hour_kursovaya_work_or_project(umk_data): #возвращает количество часов отведенных на курсовой проект или работу
    res = "-/-/-"
    js = json.JSONDecoder()
    if umk_data.table_samost_hour:
        list = js.decode(umk_data.table_samost_hour)

        for item in list:
            if re.findall('Выполнение курсового проекта / курсовой работы', item[2]):
                res = num_is_empty(item[3])
    return res

def get_hour_raschetno_graph_work(umk_data):#возвращает количество часов отведенных на расчетно-графическую работу
    res = "-/-/-"
    js = json.JSONDecoder()
    if umk_data.table_samost_hour:
        list = js.decode(umk_data.table_samost_hour)

        for item in list:
            if re.findall('Выполнение расчетно-графических домашних работ', item[2]):
                res = num_is_empty(item[3])
    return res


def get_table_literature(umkdata):
    res = []
    js = json.JSONDecoder()
    if umkdata.table_literature:
        list = js.decode(umkdata.table_literature)

        for item in list:
            res.append({'main_add': item[0], 'name': item[1], 'year': item[2], 'type': item[3], 'vid': item[4],
                        'count': item[5], 'kontingent': item[6], 'obespechennost': item[7], 'place': item[8], 'electr_variant': item[9]})
    return res

def get_kontrolnya(plans): #только для заочной формы обучения
    res = ''
    if plans[1].kontrolnaya_work == plans[2].kontrolnaya_work and  ( (plans[1].trudoemkost_all - plans[1].hours_samost_work_sum) == (plans[2].trudoemkost_all - plans[2].hours_samost_work_sum)):
        res = 'Для заочной формы обучения предусмотрена контрольная работа объемом {0} часов в {1} семестре. '.format(plans[1].trudoemkost_all - plans[1].hours_samost_work_sum, plans[1].kontrolnaya_work)
    else:
        for i in range(1,3):
            srok = re.findall(r"\d\.\d|\d+", plans[i].get_training_form_display())[0]
            if plans[i].kontrolnaya_work:
                res += 'Для заочной ({0} {1}) формы обучения предусмотрена контрольная работа объемом {2} часов в {3} семестре. '.format(srok, "лет" if float(srok)>=5 else "года",
                                                                                                                                         plans[i].trudoemkost_all - plans[i].hours_samost_work_sum,
                                                                                                                                         plans[i].kontrolnaya_work)

    return res

def get_table_rating_day(doc_tpl, plan, umkdata):#Оценка результатов освоения учебной дисциплины для очной формы
    js = json.JSONDecoder()
    datafortable = []
    tmp = []
    id = 0

    sd = doc_tpl.new_subdoc()
    if umkdata.table_rating_ochka:
        sd.add_paragraph('Рейтинговая система оценки по дисциплине «{0}» для обучающихся направления {1} очной формы обучения'.format(plan.discipline.name, plan.get_direction()))
        table1 = sd.add_table(rows=1, cols=4, style='Table Grid')
        hdr_cells = table1.rows[0].cells
        hdr_cells[0].text = '1-ый срок предоставления результатов текущего контроля'
        hdr_cells[1].text = '2-ый срок предоставления результатов текущего контроля'
        hdr_cells[2].text = '3-ый срок предоставления результатов текущего контроля'
        hdr_cells[3].text = 'Итоговый тест'

        for item in js.decode(umkdata.table_rating_ochka):
            if re.search("Итого за \d-ую аттестацию", item[1]):
                tmp.append(item[2].split("-")[1] if re.search("0-\d+", item[2]) else item[2])
                id += 1
            if id == 3:
                datafortable.append(tmp)
                tmp = []
                id = 0

        for item in datafortable:
            row_cells = table1.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
            row_cells[2].text = item[2]
            row_cells[3].text = '100'

        sd.add_paragraph(" ")
        table = sd.add_table(rows=1, cols=4, style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'Виды контрольных мероприятий'
        hdr_cells[2].text = 'Баллы'
        hdr_cells[3].text = '№ недели'

        for item in js.decode(umkdata.table_rating_ochka):
            row_cells = table.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
            row_cells[2].text = item[2]
            row_cells[3].text = item[3]

    return sd

def get_table_rating_night(doc_tpl, umkdata):#Оценка результатов освоения учебной дисциплины для заочной формы
    js = json.JSONDecoder()
    sd = doc_tpl.new_subdoc()

    if umkdata.table_rating_zaochka:
        table = sd.add_table(rows=1, cols=3, style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'Виды контрольных мероприятий'
        hdr_cells[2].text = 'Баллы'

        for item in js.decode(umkdata.table_rating_zaochka):
            row_cells = table.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
            row_cells[2].text = item[2]

    return sd